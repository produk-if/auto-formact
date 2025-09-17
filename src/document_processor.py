import os
import logging
from datetime import datetime
from docx import Document
from .validator import DocumentValidator
from .corrector import DocumentCorrector

class DocumentProcessor:
    """Main document processing coordinator"""

    def __init__(self, config):
        self.config = config
        self.validator = DocumentValidator(config)
        self.corrector = DocumentCorrector(config)

        # Set up logging
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler('temp/processing.log'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)

    def process_document(self, document_path, file_id):
        """Process a document through validation and optional correction"""
        self.logger.info(f"Starting processing for document: {document_path}")

        try:
            # Load document
            doc = Document(document_path)
            self.logger.info(f"Document loaded successfully: {len(doc.paragraphs)} paragraphs")

            # Create backup
            backup_path = self._create_backup(document_path, file_id)

            # Validate document
            self.logger.info("Starting document validation...")
            violations = self.validator.validate_document(document_path)
            self.logger.info(f"Validation complete: {len(violations)} violations found")

            # Categorize violations by severity
            errors = [v for v in violations if v['severity'] == 'error']
            warnings = [v for v in violations if v['severity'] == 'warning']
            suggestions = [v for v in violations if v['severity'] == 'suggestion']

            # Get auto-correctable violations
            auto_correctable = [v for v in violations if v.get('auto_correctable', False)]

            processing_result = {
                'timestamp': datetime.now().isoformat(),
                'document_info': self._get_document_info(doc),
                'validation': {
                    'total_violations': len(violations),
                    'errors': len(errors),
                    'warnings': len(warnings),
                    'suggestions': len(suggestions),
                    'violations': violations
                },
                'auto_correction': {
                    'available': len(auto_correctable),
                    'correctable_violations': auto_correctable
                },
                'backup_created': backup_path,
                'processing_status': 'completed'
            }

            self.logger.info("Document processing completed successfully")
            return processing_result

        except Exception as e:
            self.logger.error(f"Error processing document: {str(e)}")
            raise

    def _create_backup(self, document_path, file_id):
        """Create a backup of the original document"""
        try:
            backup_filename = f"{file_id}_backup.docx"
            backup_path = os.path.join('temp', backup_filename)

            # Copy file
            with open(document_path, 'rb') as src, open(backup_path, 'wb') as dst:
                dst.write(src.read())

            self.logger.info(f"Backup created: {backup_path}")
            return backup_path

        except Exception as e:
            self.logger.error(f"Failed to create backup: {str(e)}")
            raise

    def _get_document_info(self, doc):
        """Extract basic document information"""
        try:
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)

            # Count words (simple approach)
            word_count = 0
            for paragraph in doc.paragraphs:
                word_count += len(paragraph.text.split())

            # Get document properties
            props = doc.core_properties
            title = props.title or "Untitled"
            author = props.author or "Unknown"

            return {
                'title': title,
                'author': author,
                'paragraph_count': paragraph_count,
                'table_count': table_count,
                'estimated_word_count': word_count,
                'created': props.created.isoformat() if props.created else None,
                'modified': props.modified.isoformat() if props.modified else None
            }

        except Exception as e:
            self.logger.error(f"Error extracting document info: {str(e)}")
            return {
                'title': 'Error extracting title',
                'author': 'Error extracting author',
                'paragraph_count': 0,
                'table_count': 0,
                'estimated_word_count': 0
            }