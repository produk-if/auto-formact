import os
import logging
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from docx import Document
from .validator import DocumentValidator

class ReportGenerator:
    """Generates PDF compliance reports"""

    def __init__(self, config):
        self.config = config
        self.logger = logging.getLogger(__name__)

    def generate_pdf_report(self, document_path, file_id):
        """Generate a comprehensive PDF compliance report"""
        try:
            # Validate document to get violations
            validator = DocumentValidator(self.config)
            violations = validator.validate_document(document_path)

            # Create report filename
            report_filename = f"compliance_report_{file_id}.pdf"
            report_path = os.path.join('temp', report_filename)

            # Create PDF document
            doc_pdf = SimpleDocTemplate(report_path, pagesize=A4)
            story = []

            # Get styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=16,
                spaceAfter=30,
                alignment=1  # Center alignment
            )

            # Add title
            title_text = "LAPORAN KEPATUHAN DOKUMEN TESIS<br/>UNIVERSITAS MUHAMMADIYAH MAKASSAR"
            story.append(Paragraph(title_text, title_style))
            story.append(Spacer(1, 20))

            # Add document information
            doc_word = Document(document_path)
            doc_info = self._get_document_info(doc_word)

            info_data = [
                ['Informasi Dokumen', ''],
                ['Nama File', os.path.basename(document_path)],
                ['Tanggal Pemrosesan', datetime.now().strftime('%d %B %Y %H:%M:%S')],
                ['Judul Dokumen', doc_info.get('title', 'Tidak tersedia')],
                ['Penulis', doc_info.get('author', 'Tidak tersedia')],
                ['Jumlah Paragraf', str(doc_info.get('paragraph_count', 0))],
                ['Jumlah Tabel', str(doc_info.get('table_count', 0))],
                ['Estimasi Jumlah Kata', str(doc_info.get('estimated_word_count', 0))]
            ]

            info_table = Table(info_data, colWidths=[3*inch, 3*inch])
            info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(info_table)
            story.append(Spacer(1, 20))

            # Add summary
            severity_summary = validator.get_severity_summary(violations)
            story.append(Paragraph("RINGKASAN VALIDASI", styles['Heading2']))

            summary_data = [
                ['Tingkat Pelanggaran', 'Jumlah'],
                ['Error (Kesalahan Kritis)', str(severity_summary.get('error', 0))],
                ['Warning (Peringatan)', str(severity_summary.get('warning', 0))],
                ['Suggestion (Saran)', str(severity_summary.get('suggestion', 0))],
                ['Total Pelanggaran', str(len(violations))]
            ]

            summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (0, 1), colors.lightcoral),  # Error row
                ('BACKGROUND', (0, 2), (0, 2), colors.lightyellow),  # Warning row
                ('BACKGROUND', (0, 3), (0, 3), colors.lightblue),   # Suggestion row
                ('BACKGROUND', (0, 4), (0, 4), colors.lightgrey),   # Total row
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(summary_table)
            story.append(Spacer(1, 20))

            # Add detailed violations
            if violations:
                story.append(Paragraph("DETAIL PELANGGARAN", styles['Heading2']))

                # Group violations by severity
                errors = [v for v in violations if v.get('severity') == 'error']
                warnings = [v for v in violations if v.get('severity') == 'warning']
                suggestions = [v for v in violations if v.get('severity') == 'suggestion']

                if errors:
                    story.append(Paragraph("KESALAHAN KRITIS (ERROR)", styles['Heading3']))
                    for i, error in enumerate(errors, 1):
                        error_text = f"{i}. {error.get('message', 'Tidak ada pesan')}"
                        if error.get('location'):
                            error_text += f" (Lokasi: {error['location']})"
                        story.append(Paragraph(error_text, styles['Normal']))
                    story.append(Spacer(1, 12))

                if warnings:
                    story.append(Paragraph("PERINGATAN (WARNING)", styles['Heading3']))
                    for i, warning in enumerate(warnings, 1):
                        warning_text = f"{i}. {warning.get('message', 'Tidak ada pesan')}"
                        if warning.get('location'):
                            warning_text += f" (Lokasi: {warning['location']})"
                        story.append(Paragraph(warning_text, styles['Normal']))
                    story.append(Spacer(1, 12))

                if suggestions:
                    story.append(Paragraph("SARAN PERBAIKAN (SUGGESTION)", styles['Heading3']))
                    for i, suggestion in enumerate(suggestions, 1):
                        suggestion_text = f"{i}. {suggestion.get('message', 'Tidak ada pesan')}"
                        if suggestion.get('location'):
                            suggestion_text += f" (Lokasi: {suggestion['location']})"
                        story.append(Paragraph(suggestion_text, styles['Normal']))
                    story.append(Spacer(1, 12))

            else:
                story.append(Paragraph("TIDAK ADA PELANGGARAN DITEMUKAN", styles['Heading2']))
                story.append(Paragraph("Dokumen ini telah memenuhi semua aturan formatting UNISMUH Makassar.", styles['Normal']))

            # Add guidelines reference
            story.append(Spacer(1, 20))
            story.append(Paragraph("REFERENSI PEDOMAN", styles['Heading2']))
            guidelines_text = f"""
            Validasi ini berdasarkan pedoman resmi:
            <br/>• {self.config['university']['guidelines']}
            <br/>• Universitas: {self.config['university']['name']} ({self.config['university']['abbreviation']})
            <br/>• Tanggal validasi: {datetime.now().strftime('%d %B %Y')}
            """
            story.append(Paragraph(guidelines_text, styles['Normal']))

            # Add footer
            story.append(Spacer(1, 30))
            footer_text = "Laporan ini dihasilkan secara otomatis oleh Sistem Formatting Tesis UNISMUH"
            story.append(Paragraph(footer_text, styles['Normal']))

            # Build PDF
            doc_pdf.build(story)

            self.logger.info(f"PDF report generated: {report_path}")
            return report_path

        except Exception as e:
            self.logger.error(f"Error generating PDF report: {str(e)}")
            raise

    def _get_document_info(self, doc):
        """Extract basic document information"""
        try:
            paragraph_count = len(doc.paragraphs)
            table_count = len(doc.tables)

            # Count words
            word_count = 0
            for paragraph in doc.paragraphs:
                word_count += len(paragraph.text.split())

            # Get document properties
            props = doc.core_properties
            title = props.title or "Tidak tersedia"
            author = props.author or "Tidak tersedia"

            return {
                'title': title,
                'author': author,
                'paragraph_count': paragraph_count,
                'table_count': table_count,
                'estimated_word_count': word_count,
                'created': props.created,
                'modified': props.modified
            }

        except Exception as e:
            self.logger.error(f"Error extracting document info: {str(e)}")
            return {
                'title': 'Error',
                'author': 'Error',
                'paragraph_count': 0,
                'table_count': 0,
                'estimated_word_count': 0
            }

    def generate_summary_report(self, violations):
        """Generate a simple text summary report"""
        try:
            summary = []
            summary.append("=== RINGKASAN VALIDASI DOKUMEN ===")
            summary.append(f"Tanggal: {datetime.now().strftime('%d %B %Y %H:%M:%S')}")
            summary.append("")

            # Count by severity
            severity_counts = {'error': 0, 'warning': 0, 'suggestion': 0}
            for violation in violations:
                severity = violation.get('severity', 'error')
                if severity in severity_counts:
                    severity_counts[severity] += 1

            summary.append(f"Total pelanggaran: {len(violations)}")
            summary.append(f"- Error: {severity_counts['error']}")
            summary.append(f"- Warning: {severity_counts['warning']}")
            summary.append(f"- Suggestion: {severity_counts['suggestion']}")
            summary.append("")

            if violations:
                summary.append("=== DETAIL PELANGGARAN ===")
                for i, violation in enumerate(violations, 1):
                    severity = violation.get('severity', 'error').upper()
                    message = violation.get('message', 'Tidak ada pesan')
                    location = violation.get('location', 'Tidak diketahui')
                    summary.append(f"{i}. [{severity}] {message} (Lokasi: {location})")
            else:
                summary.append("Tidak ada pelanggaran ditemukan.")

            return "\n".join(summary)

        except Exception as e:
            self.logger.error(f"Error generating summary report: {str(e)}")
            return f"Error generating report: {str(e)}"