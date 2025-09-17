import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .document_restructurer import DocumentRestructurer

class DocumentCorrector:
    """Applies automatic corrections to documents"""

    def __init__(self, config):
        self.config = config
        self.logger = logging.getLogger(__name__)
        self.restructurer = DocumentRestructurer(config)

    def apply_corrections(self, document_path, corrections_to_apply):
        """Apply specified corrections to document"""
        corrections_applied = []
        corrections_failed = []

        try:
            doc = Document(document_path)

            # Group corrections by type for efficient processing
            corrections_by_type = {}
            for correction in corrections_to_apply:
                correction_type = correction.get('type')
                if correction_type not in corrections_by_type:
                    corrections_by_type[correction_type] = []
                corrections_by_type[correction_type].append(correction)

            # Apply margin corrections
            if 'margin' in corrections_by_type:
                result = self._apply_margin_corrections(doc, corrections_by_type['margin'])
                corrections_applied.extend(result['applied'])
                corrections_failed.extend(result['failed'])

            # Apply font corrections
            if 'font' in corrections_by_type:
                result = self._apply_font_corrections(doc, corrections_by_type['font'])
                corrections_applied.extend(result['applied'])
                corrections_failed.extend(result['failed'])

            # Apply font size corrections
            if 'font_size' in corrections_by_type:
                result = self._apply_font_size_corrections(doc, corrections_by_type['font_size'])
                corrections_applied.extend(result['applied'])
                corrections_failed.extend(result['failed'])

            # Apply line spacing corrections
            if 'line_spacing' in corrections_by_type:
                result = self._apply_line_spacing_corrections(doc, corrections_by_type['line_spacing'])
                corrections_applied.extend(result['applied'])
                corrections_failed.extend(result['failed'])

            # Apply heading alignment corrections
            if 'heading_alignment' in corrections_by_type:
                result = self._apply_heading_alignment_corrections(doc, corrections_by_type['heading_alignment'])
                corrections_applied.extend(result['applied'])
                corrections_failed.extend(result['failed'])

            # Apply text formatting corrections
            if 'decimal_separator' in corrections_by_type:
                result = self._apply_decimal_separator_corrections(doc, corrections_by_type['decimal_separator'])
                corrections_applied.extend(result['applied'])
                corrections_failed.extend(result['failed'])

            # Apply document restructuring
            if 'document_restructure' in corrections_by_type:
                result = self._apply_document_restructuring(document_path, corrections_by_type['document_restructure'])
                corrections_applied.extend(result['applied'])
                corrections_failed.extend(result['failed'])
                # If restructuring was successful, use the restructured file path
                if result.get('restructured_file_path'):
                    document_path = result['restructured_file_path']
                    doc = Document(document_path)  # Reload document

            # Save corrected document
            corrected_path = self._save_corrected_document(doc, document_path)

            return {
                'applied': corrections_applied,
                'failed': corrections_failed,
                'corrected_file_path': corrected_path
            }

        except Exception as e:
            self.logger.error(f"Error applying corrections: {str(e)}")
            corrections_failed.append({
                'type': 'system_error',
                'message': f'Failed to apply corrections: {str(e)}'
            })
            return {
                'applied': corrections_applied,
                'failed': corrections_failed,
                'corrected_file_path': None
            }

    def _apply_margin_corrections(self, doc, margin_corrections):
        """Apply margin corrections to all sections"""
        applied = []
        failed = []

        try:
            for section in doc.sections:
                for correction in margin_corrections:
                    margin_type = correction.get('margin')
                    value_cm = correction.get('value')

                    if margin_type == 'top':
                        section.top_margin = Cm(value_cm)
                        applied.append(f'Set top margin to {value_cm}cm')
                    elif margin_type == 'bottom':
                        section.bottom_margin = Cm(value_cm)
                        applied.append(f'Set bottom margin to {value_cm}cm')
                    elif margin_type == 'left':
                        section.left_margin = Cm(value_cm)
                        applied.append(f'Set left margin to {value_cm}cm')
                    elif margin_type == 'right':
                        section.right_margin = Cm(value_cm)
                        applied.append(f'Set right margin to {value_cm}cm')

        except Exception as e:
            failed.append(f'Failed to apply margin corrections: {str(e)}')

        return {'applied': applied, 'failed': failed}

    def _apply_font_corrections(self, doc, font_corrections):
        """Apply font family corrections"""
        applied = []
        failed = []

        try:
            expected_font = self.config['typography']['body_font']['family']

            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.font.name and run.font.name != expected_font:
                        old_font = run.font.name
                        run.font.name = expected_font
                        applied.append(f'Changed font from {old_font} to {expected_font}')

        except Exception as e:
            failed.append(f'Failed to apply font corrections: {str(e)}')

        return {'applied': applied, 'failed': failed}

    def _apply_font_size_corrections(self, doc, font_size_corrections):
        """Apply font size corrections"""
        applied = []
        failed = []

        try:
            expected_size = int(self.config['typography']['body_font']['size'].replace('pt', ''))

            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.font.size and run.font.size.pt != expected_size:
                        old_size = run.font.size.pt
                        run.font.size = Pt(expected_size)
                        applied.append(f'Changed font size from {old_size}pt to {expected_size}pt')

        except Exception as e:
            failed.append(f'Failed to apply font size corrections: {str(e)}')

        return {'applied': applied, 'failed': failed}

    def _apply_line_spacing_corrections(self, doc, spacing_corrections):
        """Apply line spacing corrections"""
        applied = []
        failed = []

        try:
            expected_spacing = self.config['typography']['line_spacing']['body_text']

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only apply to non-empty paragraphs
                    current_spacing = paragraph.paragraph_format.line_spacing
                    if current_spacing and abs(current_spacing - expected_spacing) > 0.1:
                        paragraph.paragraph_format.line_spacing = expected_spacing
                        applied.append(f'Set line spacing to {expected_spacing}')

        except Exception as e:
            failed.append(f'Failed to apply line spacing corrections: {str(e)}')

        return {'applied': applied, 'failed': failed}

    def _apply_heading_alignment_corrections(self, doc, alignment_corrections):
        """Apply heading alignment corrections"""
        applied = []
        failed = []

        try:
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip().upper()
                if text.startswith('BAB '):
                    if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        applied.append(f'Center-aligned chapter heading: {text}')

                    # Make heading bold
                    for run in paragraph.runs:
                        if not run.bold:
                            run.bold = True
                            applied.append(f'Made chapter heading bold: {text}')

        except Exception as e:
            failed.append(f'Failed to apply heading alignment corrections: {str(e)}')

        return {'applied': applied, 'failed': failed}

    def _apply_decimal_separator_corrections(self, doc, decimal_corrections):
        """Apply decimal separator corrections"""
        applied = []
        failed = []

        try:
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                # Replace decimal dots with commas (simple approach)
                corrected_text = self._fix_decimal_separators(original_text)

                if corrected_text != original_text:
                    # Clear existing runs and add corrected text
                    paragraph.clear()
                    paragraph.add_run(corrected_text)
                    applied.append(f'Fixed decimal separators in paragraph')

        except Exception as e:
            failed.append(f'Failed to apply decimal separator corrections: {str(e)}')

        return {'applied': applied, 'failed': failed}

    def _fix_decimal_separators(self, text):
        """Fix decimal separators in text"""
        import re
        # Replace numbers like 50.5 with 50,5 (but not version numbers or ranges)
        pattern = r'\b(\d+)\.(\d{1,3})\b(?!\d)'
        return re.sub(pattern, r'\1,\2', text)

    def _save_corrected_document(self, doc, original_path):
        """Save the corrected document"""
        try:
            # Generate corrected filename
            base_name = os.path.basename(original_path)
            name_without_ext = os.path.splitext(base_name)[0]

            # Extract file_id from filename (assuming format: file_id_original_name.docx)
            parts = name_without_ext.split('_', 1)
            if len(parts) >= 2:
                file_id = parts[0]
                corrected_filename = f"{file_id}_corrected.docx"
            else:
                corrected_filename = f"corrected_{base_name}"

            corrected_path = os.path.join('temp', corrected_filename)

            # Save the document
            doc.save(corrected_path)

            self.logger.info(f"Corrected document saved: {corrected_path}")
            return corrected_path

        except Exception as e:
            self.logger.error(f"Failed to save corrected document: {str(e)}")
            raise

    def get_auto_correctable_violations(self, violations):
        """Filter violations to get only auto-correctable ones"""
        return [v for v in violations if v.get('auto_correctable', False)]

    def apply_all_auto_corrections(self, document_path):
        """Apply all available auto-corrections"""
        try:
            # This would be called when user wants to auto-fix everything
            from .validator import DocumentValidator

            validator = DocumentValidator(self.config)
            violations = validator.validate_document(document_path)

            auto_correctable = self.get_auto_correctable_violations(violations)

            if auto_correctable:
                # Extract corrections from violations
                corrections_to_apply = []
                for violation in auto_correctable:
                    if 'correction' in violation:
                        corrections_to_apply.append(violation['correction'])

                return self.apply_corrections(document_path, corrections_to_apply)
            else:
                return {
                    'applied': [],
                    'failed': [],
                    'corrected_file_path': None,
                    'message': 'No auto-correctable violations found'
                }

        except Exception as e:
            self.logger.error(f"Error in auto-correction: {str(e)}")
            return {
                'applied': [],
                'failed': [f'Auto-correction failed: {str(e)}'],
                'corrected_file_path': None
            }

    def _apply_document_restructuring(self, document_path, restructure_corrections):
        """Apply document restructuring corrections"""
        applied = []
        failed = []

        try:
            # Use the restructurer to reorganize the document
            result = self.restructurer.restructure_document(
                document_path,
                {'reorder_chapters': True}
            )

            if result['success']:
                applied.extend(result['changes_applied'])
                applied.append(f"Document restructured from order: {result.get('original_order', [])} to: {result.get('corrected_order', [])}")

                return {
                    'applied': applied,
                    'failed': failed,
                    'restructured_file_path': result['restructured_file_path']
                }
            else:
                failed.append(f"Document restructuring failed: {result.get('message', 'Unknown error')}")

        except Exception as e:
            failed.append(f'Failed to apply document restructuring: {str(e)}')

        return {
            'applied': applied,
            'failed': failed,
            'restructured_file_path': None
        }