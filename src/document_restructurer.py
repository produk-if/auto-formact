import re
import os
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentRestructurer:
    """Automatically restructures documents to match UNISMUH standards"""

    def __init__(self, config):
        self.config = config
        self.logger = logging.getLogger(__name__)

    def analyze_document_structure(self, document_path):
        """Analyze current document structure and identify issues"""
        try:
            doc = Document(document_path)
            analysis = {
                'chapters': [],
                'structure_issues': [],
                'reordering_needed': False,
                'missing_sections': [],
                'extra_sections': []
            }

            # Extract all chapters and their positions
            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip().upper()

                # Detect chapter headings (BAB I, BAB II, etc.)
                chapter_match = re.match(r'BAB\s+([IVX]+)\s+(.+)', text)
                if chapter_match:
                    roman_num = chapter_match.group(1)
                    chapter_title = chapter_match.group(2).strip()

                    # Convert roman to integer for ordering
                    chapter_num = self._roman_to_int(roman_num)

                    chapter_info = {
                        'paragraph_index': para_idx,
                        'roman_numeral': roman_num,
                        'chapter_number': chapter_num,
                        'title': chapter_title,
                        'full_text': text,
                        'subsections': []
                    }

                    analysis['chapters'].append(chapter_info)

            # Check if chapters are in correct order
            if len(analysis['chapters']) > 1:
                chapter_numbers = [ch['chapter_number'] for ch in analysis['chapters']]
                sorted_numbers = sorted(chapter_numbers)

                if chapter_numbers != sorted_numbers:
                    analysis['reordering_needed'] = True
                    analysis['structure_issues'].append({
                        'type': 'chapter_order',
                        'severity': 'error',
                        'message': f'Chapters are not in correct order: found {chapter_numbers}, should be {sorted_numbers}',
                        'auto_correctable': True
                    })

            # Check for required chapters (BAB I, II, III for proposal)
            required_chapters = self.config['document_types']['proposal']['required_sections']
            found_chapter_titles = [ch['title'] for ch in analysis['chapters']]

            for required in required_chapters:
                required_clean = required.replace('BAB I ', '').replace('BAB II ', '').replace('BAB III ', '')
                found = any(required_clean in title for title in found_chapter_titles)

                if not found:
                    analysis['missing_sections'].append({
                        'type': 'missing_chapter',
                        'severity': 'error',
                        'chapter': required,
                        'message': f'Missing required chapter: {required}',
                        'auto_correctable': False
                    })

            # Extract subsections for each chapter
            self._extract_subsections(doc, analysis)

            self.logger.info(f"Document structure analysis completed: {len(analysis['chapters'])} chapters found")
            return analysis

        except Exception as e:
            self.logger.error(f"Error analyzing document structure: {str(e)}")
            return {
                'chapters': [],
                'structure_issues': [{'type': 'analysis_error', 'message': str(e)}],
                'reordering_needed': False
            }

    def _extract_subsections(self, doc, analysis):
        """Extract subsections for each chapter"""
        try:
            current_chapter_idx = -1

            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()

                # Check if this is a chapter heading
                if re.match(r'BAB\s+[IVX]+', text.upper()):
                    current_chapter_idx += 1
                    continue

                # Check for subsections (1.1, 1.2, etc. or A., B., etc.)
                subsection_match = re.match(r'(\d+\.\d+|\d+\.\d+\.\d+|[A-Z]\.|[0-9]+\.)\s+(.+)', text)
                if subsection_match and current_chapter_idx >= 0 and current_chapter_idx < len(analysis['chapters']):
                    subsection_num = subsection_match.group(1)
                    subsection_title = subsection_match.group(2)

                    analysis['chapters'][current_chapter_idx]['subsections'].append({
                        'paragraph_index': para_idx,
                        'number': subsection_num,
                        'title': subsection_title,
                        'full_text': text
                    })

        except Exception as e:
            self.logger.error(f"Error extracting subsections: {str(e)}")

    def restructure_document(self, document_path, restructure_options):
        """Restructure document according to UNISMUH standards"""
        try:
            doc = Document(document_path)

            # Analyze current structure
            analysis = self.analyze_document_structure(document_path)

            if not analysis['reordering_needed']:
                return {
                    'success': True,
                    'message': 'Document structure is already correct',
                    'changes_applied': [],
                    'restructured_file_path': None
                }

            # Create restructured version
            restructured_doc = Document()
            changes_applied = []

            # Copy document properties
            restructured_doc.core_properties.title = doc.core_properties.title
            restructured_doc.core_properties.author = doc.core_properties.author

            # Sort chapters by correct order
            sorted_chapters = sorted(analysis['chapters'], key=lambda x: x['chapter_number'])

            # Process each chapter in correct order
            for chapter_info in sorted_chapters:
                # Add chapter heading with correct formatting
                corrected_chapter = self._create_corrected_chapter_heading(
                    restructured_doc, chapter_info
                )
                changes_applied.append(f"Reordered chapter: {chapter_info['title']}")

                # Add chapter content
                self._copy_chapter_content(
                    doc, restructured_doc, chapter_info, analysis
                )

            # Save restructured document
            restructured_path = self._save_restructured_document(
                restructured_doc, document_path
            )

            return {
                'success': True,
                'message': f'Document successfully restructured with {len(changes_applied)} changes',
                'changes_applied': changes_applied,
                'restructured_file_path': restructured_path,
                'original_order': [ch['title'] for ch in analysis['chapters']],
                'corrected_order': [ch['title'] for ch in sorted_chapters]
            }

        except Exception as e:
            self.logger.error(f"Error restructuring document: {str(e)}")
            return {
                'success': False,
                'message': f'Restructuring failed: {str(e)}',
                'changes_applied': [],
                'restructured_file_path': None
            }

    def _create_corrected_chapter_heading(self, doc, chapter_info):
        """Create properly formatted chapter heading"""
        try:
            # Convert chapter number back to roman numeral (ensure correct format)
            correct_roman = self._int_to_roman(chapter_info['chapter_number'])

            # Create chapter heading paragraph
            chapter_para = doc.add_paragraph()
            chapter_run = chapter_para.add_run(f"BAB {correct_roman} {chapter_info['title']}")

            # Apply correct formatting
            chapter_run.font.name = 'Times New Roman'
            chapter_run.font.size = Pt(12)
            chapter_run.bold = True
            chapter_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            return chapter_para

        except Exception as e:
            self.logger.error(f"Error creating chapter heading: {str(e)}")
            return None

    def _copy_chapter_content(self, source_doc, target_doc, chapter_info, analysis):
        """Copy content from source chapter to target document"""
        try:
            start_idx = chapter_info['paragraph_index']

            # Find end index (start of next chapter or end of document)
            end_idx = len(source_doc.paragraphs)
            for other_chapter in analysis['chapters']:
                if (other_chapter['paragraph_index'] > start_idx and
                    other_chapter['paragraph_index'] < end_idx):
                    end_idx = other_chapter['paragraph_index']

            # Copy paragraphs (skip the chapter heading itself)
            for i in range(start_idx + 1, end_idx):
                if i < len(source_doc.paragraphs):
                    para = source_doc.paragraphs[i]

                    # Skip empty paragraphs
                    if not para.text.strip():
                        continue

                    # Create new paragraph in target document
                    new_para = target_doc.add_paragraph()

                    # Copy runs with formatting
                    for run in para.runs:
                        new_run = new_para.add_run(run.text)

                        # Copy run formatting
                        if run.font.name:
                            new_run.font.name = run.font.name
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.bold:
                            new_run.bold = run.bold
                        if run.italic:
                            new_run.italic = run.italic

                    # Copy paragraph formatting
                    new_para.alignment = para.alignment

            # Renumber subsections if needed
            self._renumber_subsections(target_doc, chapter_info['chapter_number'])

        except Exception as e:
            self.logger.error(f"Error copying chapter content: {str(e)}")

    def _renumber_subsections(self, doc, chapter_number):
        """Renumber subsections to match chapter number"""
        try:
            subsection_counter = 1

            # Find paragraphs that need renumbering
            for para in doc.paragraphs:
                text = para.text.strip()

                # Match subsection patterns (1.1, 1.2, etc.)
                subsection_match = re.match(r'(\d+)\.(\d+)\s+(.+)', text)
                if subsection_match:
                    current_chapter = int(subsection_match.group(1))
                    subsection_title = subsection_match.group(3)

                    # Update numbering to match correct chapter
                    if current_chapter != chapter_number:
                        new_text = f"{chapter_number}.{subsection_counter} {subsection_title}"

                        # Clear paragraph and add corrected text
                        para.clear()
                        run = para.add_run(new_text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

                        subsection_counter += 1

        except Exception as e:
            self.logger.error(f"Error renumbering subsections: {str(e)}")

    def _save_restructured_document(self, doc, original_path):
        """Save the restructured document"""
        try:
            # Generate restructured filename
            base_name = os.path.basename(original_path)
            name_without_ext = os.path.splitext(base_name)[0]

            # Extract file_id if it exists
            parts = name_without_ext.split('_', 1)
            if len(parts) >= 2:
                file_id = parts[0]
                restructured_filename = f"{file_id}_restructured.docx"
            else:
                restructured_filename = f"restructured_{base_name}"

            restructured_path = os.path.join('temp', restructured_filename)

            # Save document
            doc.save(restructured_path)

            self.logger.info(f"Restructured document saved: {restructured_path}")
            return restructured_path

        except Exception as e:
            self.logger.error(f"Failed to save restructured document: {str(e)}")
            raise

    def _roman_to_int(self, roman):
        """Convert roman numeral to integer"""
        roman_numerals = {'I': 1, 'V': 5, 'X': 10}
        result = 0
        prev_value = 0

        for char in reversed(roman):
            value = roman_numerals.get(char, 0)
            if value < prev_value:
                result -= value
            else:
                result += value
            prev_value = value

        return result

    def _int_to_roman(self, num):
        """Convert integer to roman numeral"""
        values = [10, 9, 5, 4, 1]
        literals = ['X', 'IX', 'V', 'IV', 'I']

        result = ""
        for i, value in enumerate(values):
            count = num // value
            if count:
                result += literals[i] * count
                num -= value * count

        return result

    def get_restructuring_preview(self, document_path):
        """Generate a preview of restructuring changes"""
        try:
            analysis = self.analyze_document_structure(document_path)

            if not analysis['reordering_needed']:
                return {
                    'preview_available': False,
                    'message': 'No restructuring needed'
                }

            # Generate preview
            current_order = []
            corrected_order = []

            for chapter in analysis['chapters']:
                current_order.append({
                    'roman': chapter['roman_numeral'],
                    'title': chapter['title'],
                    'number': chapter['chapter_number']
                })

            sorted_chapters = sorted(analysis['chapters'], key=lambda x: x['chapter_number'])
            for chapter in sorted_chapters:
                corrected_order.append({
                    'roman': self._int_to_roman(chapter['chapter_number']),
                    'title': chapter['title'],
                    'number': chapter['chapter_number']
                })

            return {
                'preview_available': True,
                'current_order': current_order,
                'corrected_order': corrected_order,
                'changes_needed': len([ch for ch in analysis['chapters']]),
                'structure_issues': analysis['structure_issues']
            }

        except Exception as e:
            self.logger.error(f"Error generating restructuring preview: {str(e)}")
            return {
                'preview_available': False,
                'message': f'Preview generation failed: {str(e)}'
            }