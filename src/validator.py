import re
import logging
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .document_restructurer import DocumentRestructurer

class DocumentValidator:
    """Validates documents against UNISMUH formatting rules"""

    def __init__(self, config):
        self.config = config
        self.rules = config['validation_rules']
        self.logger = logging.getLogger(__name__)
        self.restructurer = DocumentRestructurer(config)

    def validate_document(self, document_path):
        """Validate entire document and return list of violations"""
        violations = []

        try:
            doc = Document(document_path)

            # Validate page setup
            violations.extend(self._validate_page_setup(doc))

            # Validate typography
            violations.extend(self._validate_typography(doc))

            # Validate document structure
            violations.extend(self._validate_structure(doc))

            # Check for structural issues that need restructuring
            violations.extend(self._validate_document_order(document_path))

            # Validate headings
            violations.extend(self._validate_headings(doc))

            # Validate tables and figures
            violations.extend(self._validate_tables_figures(doc))

            # Validate text formatting
            violations.extend(self._validate_text_formatting(doc))

            self.logger.info(f"Document validation completed: {len(violations)} violations found")
            return violations

        except Exception as e:
            self.logger.error(f"Error during validation: {str(e)}")
            return [{'type': 'system_error', 'message': f'Validation error: {str(e)}', 'severity': 'error'}]

    def _validate_page_setup(self, doc):
        """Validate page margins, size, and orientation"""
        violations = []

        try:
            sections = doc.sections
            for section_idx, section in enumerate(sections):
                # Check margins
                expected_margins = self.config['page_setup']['margins']

                top_margin_cm = section.top_margin.cm
                bottom_margin_cm = section.bottom_margin.cm
                left_margin_cm = section.left_margin.cm
                right_margin_cm = section.right_margin.cm

                # Convert expected margins (remove 'cm' suffix)
                expected_top = float(expected_margins['top'].replace('cm', ''))
                expected_bottom = float(expected_margins['bottom'].replace('cm', ''))
                expected_left = float(expected_margins['left'].replace('cm', ''))
                expected_right = float(expected_margins['right'].replace('cm', ''))

                tolerance = 0.2  # Allow 2mm tolerance

                if abs(top_margin_cm - expected_top) > tolerance:
                    violations.append({
                        'type': 'margin_error',
                        'severity': 'error',
                        'message': f'Top margin is {top_margin_cm:.1f}cm but should be {expected_top}cm according to UNISMUH guidelines',
                        'location': f'Section {section_idx + 1}',
                        'auto_correctable': True,
                        'correction': {'type': 'margin', 'margin': 'top', 'value': expected_top}
                    })

                if abs(bottom_margin_cm - expected_bottom) > tolerance:
                    violations.append({
                        'type': 'margin_error',
                        'severity': 'error',
                        'message': f'Bottom margin is {bottom_margin_cm:.1f}cm but should be {expected_bottom}cm according to UNISMUH guidelines',
                        'location': f'Section {section_idx + 1}',
                        'auto_correctable': True,
                        'correction': {'type': 'margin', 'margin': 'bottom', 'value': expected_bottom}
                    })

                if abs(left_margin_cm - expected_left) > tolerance:
                    violations.append({
                        'type': 'margin_error',
                        'severity': 'error',
                        'message': f'Left margin is {left_margin_cm:.1f}cm but should be {expected_left}cm according to UNISMUH guidelines',
                        'location': f'Section {section_idx + 1}',
                        'auto_correctable': True,
                        'correction': {'type': 'margin', 'margin': 'left', 'value': expected_left}
                    })

                if abs(right_margin_cm - expected_right) > tolerance:
                    violations.append({
                        'type': 'margin_error',
                        'severity': 'error',
                        'message': f'Right margin is {right_margin_cm:.1f}cm but should be {expected_right}cm according to UNISMUH guidelines',
                        'location': f'Section {section_idx + 1}',
                        'auto_correctable': True,
                        'correction': {'type': 'margin', 'margin': 'right', 'value': expected_right}
                    })

        except Exception as e:
            violations.append({
                'type': 'page_setup_error',
                'severity': 'error',
                'message': f'Error validating page setup: {str(e)}',
                'auto_correctable': False
            })

        return violations

    def _validate_typography(self, doc):
        """Validate font, size, and line spacing"""
        violations = []

        try:
            expected_font = self.config['typography']['body_font']['family']
            expected_size = int(self.config['typography']['body_font']['size'].replace('pt', ''))

            for para_idx, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():  # Skip empty paragraphs
                    for run in paragraph.runs:
                        if run.font.name and run.font.name != expected_font:
                            violations.append({
                                'type': 'font_error',
                                'severity': 'error',
                                'message': f'{run.font.name} font used instead of required {expected_font}',
                                'location': f'Paragraph {para_idx + 1}',
                                'auto_correctable': True,
                                'correction': {'type': 'font', 'font_name': expected_font}
                            })

                        if run.font.size and run.font.size.pt != expected_size:
                            violations.append({
                                'type': 'font_size_error',
                                'severity': 'error',
                                'message': f'Font size {run.font.size.pt}pt used instead of required {expected_size}pt',
                                'location': f'Paragraph {para_idx + 1}',
                                'auto_correctable': True,
                                'correction': {'type': 'font_size', 'size': expected_size}
                            })

                    # Check line spacing
                    if paragraph.paragraph_format.line_spacing:
                        current_spacing = paragraph.paragraph_format.line_spacing
                        expected_spacing = self.config['typography']['line_spacing']['body_text']

                        if abs(current_spacing - expected_spacing) > 0.1:
                            violations.append({
                                'type': 'line_spacing_error',
                                'severity': 'error',
                                'message': f'{current_spacing} spacing used instead of required {expected_spacing} line spacing',
                                'location': f'Paragraph {para_idx + 1}',
                                'auto_correctable': True,
                                'correction': {'type': 'line_spacing', 'spacing': expected_spacing}
                            })

        except Exception as e:
            violations.append({
                'type': 'typography_error',
                'severity': 'error',
                'message': f'Error validating typography: {str(e)}',
                'auto_correctable': False
            })

        return violations

    def _validate_structure(self, doc):
        """Validate document structure and required sections"""
        violations = []

        try:
            # Extract chapter headings
            chapter_headings = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip().upper()
                if text.startswith('BAB '):
                    chapter_headings.append(text)

            # Check required sections
            required_sections = self.config['document_types']['proposal']['required_sections']

            for required_section in required_sections:
                found = False
                for heading in chapter_headings:
                    if required_section.upper() in heading:
                        found = True
                        break

                if not found:
                    violations.append({
                        'type': 'structure_error',
                        'severity': 'error',
                        'message': f'Missing required section \'{required_section}\'',
                        'location': 'Document structure',
                        'auto_correctable': False
                    })

            # Check subsections for each chapter
            for chapter, subsections in self.config['document_types']['proposal']['subsections'].items():
                chapter_found = any(chapter.upper() in heading for heading in chapter_headings)

                if chapter_found:
                    for subsection in subsections:
                        subsection_found = any(subsection in paragraph.text
                                             for paragraph in doc.paragraphs)
                        if not subsection_found:
                            violations.append({
                                'type': 'subsection_missing',
                                'severity': 'warning',
                                'message': f'Missing subsection \'{subsection}\' in {chapter}',
                                'location': chapter,
                                'auto_correctable': False
                            })

        except Exception as e:
            violations.append({
                'type': 'structure_validation_error',
                'severity': 'error',
                'message': f'Error validating structure: {str(e)}',
                'auto_correctable': False
            })

        return violations

    def _validate_headings(self, doc):
        """Validate heading formats and styles"""
        violations = []

        try:
            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()

                # Check chapter headings
                if text.upper().startswith('BAB '):
                    # Should be centered and bold
                    if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                        violations.append({
                            'type': 'heading_alignment',
                            'severity': 'error',
                            'message': f'Chapter heading "{text}" should be center-aligned',
                            'location': f'Paragraph {para_idx + 1}',
                            'auto_correctable': True,
                            'correction': {'type': 'heading_alignment', 'alignment': 'center'}
                        })

                    # Check if bold (check first run)
                    if paragraph.runs and not paragraph.runs[0].bold:
                        violations.append({
                            'type': 'heading_bold',
                            'severity': 'error',
                            'message': f'Chapter heading "{text}" should be bold',
                            'location': f'Paragraph {para_idx + 1}',
                            'auto_correctable': True,
                            'correction': {'type': 'heading_bold', 'make_bold': True}
                        })

                    # Check format (BAB [ROMAN] [TITLE])
                    if not re.match(r'BAB [IVX]+ ', text.upper()):
                        violations.append({
                            'type': 'heading_format',
                            'severity': 'warning',
                            'message': f'Chapter heading format should be "BAB [ROMAN] [TITLE]": "{text}"',
                            'location': f'Paragraph {para_idx + 1}',
                            'auto_correctable': False
                        })

        except Exception as e:
            violations.append({
                'type': 'heading_validation_error',
                'severity': 'error',
                'message': f'Error validating headings: {str(e)}',
                'auto_correctable': False
            })

        return violations

    def _validate_tables_figures(self, doc):
        """Validate table and figure formatting"""
        violations = []

        try:
            # Validate tables
            for table_idx, table in enumerate(doc.tables):
                # Check if table has a title (look at paragraph before table)
                # This is simplified - in practice you'd need more sophisticated detection
                violations.append({
                    'type': 'table_title_check',
                    'severity': 'suggestion',
                    'message': f'Verify that Table {table_idx + 1} has proper title format "Tabel [NUMBER]. [Title]"',
                    'location': f'Table {table_idx + 1}',
                    'auto_correctable': False
                })

        except Exception as e:
            violations.append({
                'type': 'table_figure_validation_error',
                'severity': 'error',
                'message': f'Error validating tables/figures: {str(e)}',
                'auto_correctable': False
            })

        return violations

    def _validate_text_formatting(self, doc):
        """Validate text-specific formatting rules"""
        violations = []

        try:
            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text

                # Check for numbers at start of sentences
                sentences = text.split('. ')
                for sentence_idx, sentence in enumerate(sentences):
                    sentence = sentence.strip()
                    if re.match(r'^\d+', sentence):
                        violations.append({
                            'type': 'number_start_sentence',
                            'severity': 'warning',
                            'message': f'Number at sentence start should be written as words: "{sentence[:20]}..."',
                            'location': f'Paragraph {para_idx + 1}, Sentence {sentence_idx + 1}',
                            'auto_correctable': False
                        })

                # Check decimal and thousand separators
                if re.search(r'\d+\.\d+', text):  # Decimal with dot
                    violations.append({
                        'type': 'decimal_separator',
                        'severity': 'warning',
                        'message': 'Use comma (,) as decimal separator, not period (.)',
                        'location': f'Paragraph {para_idx + 1}',
                        'auto_correctable': True,
                        'correction': {'type': 'decimal_separator', 'replace_dots_with_commas': True}
                    })

        except Exception as e:
            violations.append({
                'type': 'text_formatting_error',
                'severity': 'error',
                'message': f'Error validating text formatting: {str(e)}',
                'auto_correctable': False
            })

        return violations

    def _validate_document_order(self, document_path):
        """Validate document chapter order and structure using restructurer"""
        violations = []

        try:
            # Use restructurer to analyze document structure
            analysis = self.restructurer.analyze_document_structure(document_path)

            # Add structure issues to violations
            violations.extend(analysis['structure_issues'])

            # Add specific violation for reordering if needed
            if analysis['reordering_needed']:
                violations.append({
                    'type': 'document_reordering',
                    'severity': 'error',
                    'message': 'Document chapters are not in the correct order and need to be restructured',
                    'location': 'Document structure',
                    'auto_correctable': True,
                    'correction': {
                        'type': 'document_restructure',
                        'action': 'reorder_chapters',
                        'current_order': [ch['title'] for ch in analysis['chapters']],
                        'correct_order': [ch['title'] for ch in sorted(analysis['chapters'], key=lambda x: x['chapter_number'])]
                    }
                })

            # Add violations for missing sections
            violations.extend(analysis['missing_sections'])

        except Exception as e:
            violations.append({
                'type': 'document_order_validation_error',
                'severity': 'error',
                'message': f'Error validating document order: {str(e)}',
                'auto_correctable': False
            })

        return violations

    def get_severity_summary(self, violations):
        """Get summary of violations by severity"""
        summary = {'error': 0, 'warning': 0, 'suggestion': 0}

        for violation in violations:
            severity = violation.get('severity', 'error')
            if severity in summary:
                summary[severity] += 1

        return summary